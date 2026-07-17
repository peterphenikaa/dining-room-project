from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = Document()

# Tiêu đề
title = doc.add_heading('BÁO CÁO TIẾN ĐỘ DỰ ÁN QUẢN LÝ PHÒNG ĂN', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph('Ngày báo cáo: 17/07/2026')
doc.add_paragraph('Người thực hiện: Backend Team')

# Phần 1: Tổng quan
doc.add_heading('1. Tổng quan Kiến trúc & Hạ tầng', level=1)
p = doc.add_paragraph()
p.add_run('- Docker hóa thành công:').bold = True
p.add_run(' Đã thiết lập Docker Compose bao gồm 2 container (MySQL 8 và Express App). Xử lý triệt để lỗi tương thích ts-node trên môi trường Node.js 24.\n')
p.add_run('- Môi trường Phát triển (Dev):').bold = True
p.add_run(' Tích hợp thành công Nodemon và cơ chế Hot-Reload, tự động biên dịch TypeScript sang JavaScript ngay khi có thay đổi code, tiết kiệm thời gian khởi động lại server.\n')
p.add_run('- Cơ sở dữ liệu:').bold = True
p.add_run(' Triển khai chuẩn TypeORM Migrations thay vì tự động đồng bộ (synchronize). Toàn bộ 5 bảng đã được khởi tạo thành công thông qua script.')

# Phần 2: Thiết kế Database
doc.add_heading('2. Thiết kế Cơ sở dữ liệu (Database Design)', level=1)
doc.add_paragraph('Hoàn thiện thiết kế 5 thực thể chính với các ràng buộc khóa ngoại (Foreign Key) chặt chẽ:')
doc.add_paragraph('- DiningRoom (Phòng ăn) - Bảng Gốc\n- DiningTable (Bàn ăn)\n- DiningCabinet (Tủ phòng ăn)\n- DiningChair (Ghế ăn)\n- DiningAccessory (Phụ kiện bàn ăn)', style='List Bullet')
doc.add_paragraph('Đã áp dụng ràng buộc ON DELETE CASCADE: Khi xóa Phòng ăn, toàn bộ Bàn và Tủ bên trong tự động bị xóa bỏ để tránh rác dữ liệu dư thừa.')

# Phần 3: Kiến trúc Code & API
doc.add_heading('3. Kiến trúc Code & Phát triển API', level=1)
p2 = doc.add_paragraph()
p2.add_run('- Kiến trúc chuẩn MVC:').bold = True
p2.add_run(' Đã phân tách code thành các tầng riêng biệt: Routes (Định tuyến) -> Controllers (Điều hướng) -> Services (Xử lý logic DB). Đảm bảo khả năng bảo trì, mở rộng.\n')
p2.add_run('- Xử lý Lỗi Toàn Cục (Global Error Handling):').bold = True
p2.add_run(' Loại bỏ hoàn toàn try...catch lặp lại trong Controller. Tích hợp AppError, catchAsync và errorHandler Middleware để thống nhất định dạng JSON trả về cho mọi loại lỗi trên hệ thống.\n')
p2.add_run('- API Hoàn thành:').bold = True
p2.add_run(' Đã hoàn thiện toàn bộ các API CRUD (Tạo, Đọc, Sửa, Xóa) cho thực thể DiningRoom (Phòng ăn).')

# Phần 4: Kế hoạch tiếp theo
doc.add_heading('4. Kế hoạch tiếp theo', level=1)
doc.add_paragraph('- Triển khai các API quản lý Bàn Ăn (DiningTable) với nghiệp vụ gán Khóa ngoại vào Phòng Ăn.', style='List Bullet')
doc.add_paragraph('- Triển khai các API truy vấn phức tạp (Eager Loading) lấy thông tin Phòng + Bàn + Ghế trong 1 lần gọi.', style='List Bullet')

doc.save('Bao_Cao_Tien_Do.docx')
print("Đã tạo thành công file Bao_Cao_Tien_Do.docx!")
